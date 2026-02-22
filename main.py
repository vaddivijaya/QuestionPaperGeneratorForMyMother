import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import requests

st.set_page_config(page_title="Question Paper Maker")
st.title("📄 Question Paper Maker")

# ---------- Telugu Transliteration ----------


def english_to_telugu(text: str) -> str:
    """
    Converts English phonetic typing into Telugu script.
    Example: 'meeru ela unnaru' → 'మీరు ఎలా ఉన్నారు'
    """
    try:
        url = "https://inputtools.google.com/request"
        params = {
            "text": text,
            "itc": "te-t-i0-und",
            "num": 1
        }

        response = requests.get(url, params=params)
        result = response.json()

        if result[0] == "SUCCESS":
            return result[1][0][1][0]

        return text

    except Exception:
        return text


# ---------- Helper: force table borders ----------

def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()

            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                elem = OxmlElement(f"w:{edge}")
                elem.set(qn("w:val"), "single")
                elem.set(qn("w:sz"), "12")
                elem.set(qn("w:space"), "0")
                elem.set(qn("w:color"), "000000")
                tcBorders.append(elem)

            tcPr.append(tcBorders)


# ---------- Upload template ----------

uploaded_template = st.file_uploader(
    "Upload Question Paper Template (.docx)",
    type=["docx"]
)

# ---------- Session state ----------

if "questions" not in st.session_state:
    st.session_state.questions = []

# ---------- Question type selector ----------

st.subheader("Add Question")

q_type = st.selectbox(
    "Select Question Type",
    ["Text", "Image", "Match the Following", "Answer Table"]
)

question_data = None


# ================= TEXT =================

# ================= TEXT =================
def convert_to_telugu():
    if st.session_state.telugu_text:
        converted = "\n".join(
            english_to_telugu(line)
            for line in st.session_state.telugu_text.split("\n")
        )
        st.session_state.telugu_text = converted


if q_type == "Text":

    st.text_input("Enter question text", key="text_q")

    st.text_area("for Telugu Conversion", key="telugu_text")

    st.button("Convert to Telugu", on_click=convert_to_telugu)

    if st.session_state.text_q:
        question_data = {
            "type": "text",
            "content": st.session_state.text_q
        }


# ================= IMAGE =================

elif q_type == "Image":
    img_file = st.file_uploader("Upload image", type=["png", "jpg", "jpeg"])
    caption = st.text_input("Optional caption", key="img_caption")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Telugu Caption"):
            if caption:
                st.session_state.img_caption = english_to_telugu(caption)
                st.rerun()

    if img_file:
        question_data = {
            "type": "image",
            "image": img_file.read(),
            "caption": caption
        }


# ================= MATCH THE FOLLOWING =================

elif q_type == "Match the Following":
    left = st.text_area("Left items (one per line)", key="left_text")
    right = st.text_area("Right items (one per line)", key="right_text")

    st.text_area("for Telugu Conversion", key="telugu_text")

    st.button("Convert to Telugu", on_click=convert_to_telugu)

    if left and right:
        question_data = {
            "type": "match",
            "left": [x.strip() for x in left.split("\n") if x.strip()],
            "right": [x.strip() for x in right.split("\n") if x.strip()],
        }


# ================= ANSWER TABLE =================

elif q_type == "Answer Table":
    rows = st.number_input("Rows", 1, 10, 3)
    cols = st.number_input("Columns", 1, 10, 3)

    st.write("Enter table cell values:")

    table_values = []
    for r in range(rows):
        row_vals = []
        cols_ui = st.columns(cols)
        for c in range(cols):
            val = cols_ui[c].text_input(f"R{r+1}C{c+1}", key=f"cell_{r}_{c}")
            row_vals.append(val)
        table_values.append(row_vals)

    question_data = {
        "type": "table",
        "rows": rows,
        "cols": cols,
        "data": table_values,
    }


# ---------- Add / Clear buttons ----------

col1, col2 = st.columns(2)

with col1:
    if st.button("➕ Add Question"):
        if question_data:
            st.session_state.questions.append(question_data)
            st.rerun()
        else:
            st.warning("Please fill the question details first.")

with col2:
    if st.button("🗑 Clear All"):
        st.session_state.questions = []
        st.rerun()


# ---------- Show added questions ----------

st.subheader("Questions Added")

if not st.session_state.questions:
    st.info("No questions added yet.")
else:
    for i, q in enumerate(st.session_state.questions, 1):

        if q["type"] == "text":
            st.write(f"{i}. {q['content']}")

        elif q["type"] == "image":
            st.write(f"{i}. Image Question")
            st.image(q["image"], width=200)

        elif q["type"] == "match":
            st.write(f"{i}. Match the Following")

        elif q["type"] == "table":
            st.write(f"{i}. Table ({q['rows']} × {q['cols']})")


# ---------- Generate DOCX ----------

st.subheader("Generate Question Paper")

if st.button("📥 Preview & Download DOCX"):

    if not uploaded_template:
        st.warning("Please upload the template first.")
        st.stop()

    if not st.session_state.questions:
        st.warning("Please add at least one question.")
        st.stop()

    doc = Document(uploaded_template)
    doc.add_paragraph("\nQuestions:\n")

    for i, q in enumerate(st.session_state.questions, 1):

        if q["type"] == "text":
            doc.add_paragraph(f"{i}. {q['content']}")

        elif q["type"] == "image":
            doc.add_paragraph(f"{i}.")
            doc.add_picture(BytesIO(q["image"]), width=Inches(4))
            if q["caption"]:
                doc.add_paragraph(q["caption"])

        elif q["type"] == "match":
            doc.add_paragraph(f"{i}. Match the Following:")
            rows = max(len(q["left"]), len(q["right"]))
            table = doc.add_table(rows=rows, cols=3)

            for r in range(rows):
                if r < len(q["left"]):
                    table.rows[r].cells[0].text = q["left"][r]
                table.rows[r].cells[1].text = "[   ]"
                if r < len(q["right"]):
                    table.rows[r].cells[2].text = q["right"][r]

        elif q["type"] == "table":
            doc.add_paragraph(f"{i}.")
            table = doc.add_table(rows=q["rows"], cols=q["cols"])

            for r in range(q["rows"]):
                for c in range(q["cols"]):
                    text = q["data"][r][c] if q["data"][r][c] else " "
                    table.rows[r].cells[c].text = text

            set_table_borders(table)

        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Question paper ready!")

    st.download_button(
        "⬇ Download Final Question Paper",
        data=buffer,
        file_name="question_paper.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
